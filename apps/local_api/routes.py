"""API routes for PolicyCapture Local."""
import json
import logging
import os
import re
import shutil
import threading
from pathlib import Path

from fastapi import APIRouter, HTTPException, Request, UploadFile, File, Query
from fastapi.responses import FileResponse, HTMLResponse, StreamingResponse

from packages.shared.schemas import (
    CreateJobRequest, RegisterVideoRequest, UpdateScreenshotRequest,
    UpdateJobTitleRequest,
)
from packages.shared.database import (
    init_db, create_job, get_job, list_jobs, update_job, update_job_status, delete_job,
    create_frame, get_frame, get_frames_for_job, get_screenshots_for_job, update_screenshot, get_screenshot,
    get_sections_for_job, get_report_for_job, create_report,
    create_screenshot, create_section,
)
from packages.shared.utils import (
    generate_id, ensure_dir, get_job_dir, get_job_subdir, validate_video_path,
)
from packages.shared.config import DATA_DIR, JOBS_DIR
from packages.core.pipeline.orchestrator import PipelineOrchestrator
from packages.core.pipeline.validate_video import validate_video
from packages.core.pipeline.sample_frames import sample_frames
from packages.core.pipeline.preprocess_frame import compute_blur_score, compute_stability_score
from packages.core.pipeline.detect_relevance import detect_relevance
from packages.core.pipeline.scene_change import detect_scene_changes
from packages.core.pipeline.detect_elements import detect_elements
from packages.core.pipeline.generate_report import generate_html_report, generate_pdf_report

logger = logging.getLogger(__name__)
router = APIRouter()


def _generate_auto_title(job_id: str) -> str | None:
    """Generate a descriptive title from OCR text of the first few frames.

    Analyzes extracted text from the first 5 frames with text, picks the most
    frequent meaningful words/phrases to create a short descriptive title.
    Returns None if no meaningful title can be generated.
    """
    frames = get_frames_for_job(job_id, limit=50)
    # Get frames that have extracted text, sorted by timestamp
    text_frames = [f for f in frames if f.get("extracted_text", "").strip()]
    if not text_frames:
        return None

    # Use first 5 frames with text
    text_frames = text_frames[:5]
    all_text = " ".join(f["extracted_text"] for f in text_frames)

    # Tokenize and count meaningful words (skip short/common words)
    stopwords = {
        "the", "a", "an", "is", "are", "was", "were", "be", "been", "being",
        "have", "has", "had", "do", "does", "did", "will", "would", "could",
        "should", "may", "might", "shall", "can", "need", "dare", "ought",
        "and", "but", "or", "nor", "not", "so", "yet", "both", "either",
        "neither", "each", "every", "all", "any", "few", "more", "most",
        "other", "some", "such", "no", "only", "own", "same", "than",
        "too", "very", "just", "because", "as", "until", "while", "of",
        "at", "by", "for", "with", "about", "against", "between", "through",
        "during", "before", "after", "above", "below", "to", "from", "up",
        "down", "in", "out", "on", "off", "over", "under", "again", "further",
        "then", "once", "here", "there", "when", "where", "why", "how",
        "this", "that", "these", "those", "i", "me", "my", "myself", "we",
        "our", "ours", "you", "your", "yours", "he", "him", "his", "she",
        "her", "hers", "it", "its", "they", "them", "their", "what", "which",
        "who", "whom", "if", "else", "also", "into", "page", "click", "enter",
        "select", "please", "next", "back", "home", "menu", "file", "edit",
        "view", "help", "new", "open", "close", "save", "print", "search",
    }
    words = re.findall(r'[A-Za-z]{3,}', all_text)
    word_counts: dict[str, int] = {}
    for w in words:
        wl = w.lower()
        if wl not in stopwords and len(wl) >= 3:
            # Preserve the original casing of the most common form
            key = wl
            word_counts[key] = word_counts.get(key, 0) + 1

    if not word_counts:
        return None

    # Get top words by frequency
    sorted_words = sorted(word_counts.items(), key=lambda x: -x[1])
    top_words = [w for w, _ in sorted_words[:5]]

    # Capitalize and join
    title = " ".join(w.capitalize() for w in top_words)

    # Cap length
    if len(title) > 60:
        title = title[:57] + "..."

    return title

# Track running pipelines
_running_jobs: dict[str, threading.Thread] = {}

# UUID pattern for path validation
_UUID_RE = re.compile(r"^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$")
_VALID_ARTIFACT_TYPES = {"frames", "screenshots", "thumbnails", "reports", "input", "processed_frames", "annotated_frames"}


def _validate_job_id(job_id: str):
    """Validate job_id is a UUID to prevent path traversal."""
    if not _UUID_RE.match(job_id):
        raise HTTPException(status_code=400, detail="Invalid job ID format")


# --- Job endpoints ---

@router.post("/jobs")
def api_create_job(req: CreateJobRequest):
    job_id = generate_id()
    ensure_dir(get_job_dir(job_id))
    ensure_dir(get_job_subdir(job_id, "input"))

    source_path = ""
    if req.source_video_path:
        valid, msg = validate_video_path(req.source_video_path)
        if not valid:
            raise HTTPException(status_code=400, detail=msg)
        source_path = req.source_video_path

    job = create_job(job_id=job_id, title=req.title, source_video_path=source_path)
    return job


@router.get("/jobs")
def api_list_jobs():
    return list_jobs()


@router.get("/jobs/{job_id}")
def api_get_job(job_id: str):
    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return job


@router.delete("/jobs/{job_id}")
def api_delete_job(job_id: str):
    """Delete a job, its database records, and all files on disk."""
    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    # Remove files on disk
    job_dir = get_job_dir(job_id)
    if job_dir.exists():
        shutil.rmtree(job_dir)

    # Remove DB records
    delete_job(job_id)
    return {"status": "deleted", "id": job_id}


@router.patch("/jobs/{job_id}/title")
def api_update_job_title(job_id: str, req: UpdateJobTitleRequest):
    """Update the title of a job."""
    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    title = req.title.strip()
    if not title:
        raise HTTPException(status_code=400, detail="Title cannot be empty")
    updated = update_job(job_id, title=title)
    return updated


@router.post("/jobs/{job_id}/auto-title")
def api_auto_title(job_id: str):
    """Generate an auto-title from OCR text for this job."""
    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    auto_title = _generate_auto_title(job_id)
    if not auto_title:
        raise HTTPException(status_code=422, detail="Could not generate title - no text found in frames")
    updated = update_job(job_id, title=auto_title)
    return updated


@router.post("/jobs/{job_id}/backfill-confidence")
def api_backfill_confidence(job_id: str):
    """Recompute ocr_confidence for frames that have text but confidence=0.

    Runs OCR on each qualifying frame to populate the confidence value.
    Useful for past jobs where confidence wasn't stored.
    """
    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    frames = get_frames_for_job(job_id, limit=2000)
    need_backfill = [
        f for f in frames
        if f.get("extracted_text", "").strip() and f.get("ocr_confidence", 0) == 0
    ]

    if not need_backfill:
        return {"updated": 0, "message": "All frames already have confidence values"}

    from packages.core.pipeline.detect_elements import detect_elements
    from packages.shared.database import _get_conn
    conn = _get_conn()
    updated = 0
    for frame in need_backfill:
        img_path = frame.get("source_image_path", "")
        if not img_path or not os.path.isfile(img_path):
            continue
        try:
            result = detect_elements(img_path, quick=True)
            conf = result.get("ocr_confidence", 0)
            if conf > 0:
                conn.execute(
                    "UPDATE frames SET ocr_confidence = ? WHERE id = ?",
                    (conf, frame["id"]),
                )
                updated += 1
        except Exception as e:
            logger.warning("Backfill failed for frame %s: %s", frame["id"], e)
            continue
    conn.commit()
    return {"updated": updated, "total": len(need_backfill)}


@router.post("/jobs/{job_id}/upload")
async def api_upload_video(job_id: str, file: UploadFile = File(...)):
    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    input_dir = ensure_dir(get_job_subdir(job_id, "input"))
    # Sanitize filename
    safe_name = re.sub(r"[^\w\-.]", "_", file.filename or "recording.mp4")
    dest = input_dir / safe_name

    with open(dest, "wb") as f:
        content = await file.read()
        f.write(content)

    update_job(job_id, source_video_path=str(dest))
    return get_job(job_id)


@router.post("/jobs/{job_id}/crop")
async def api_set_crop(job_id: str, request: Request):
    """Store crop region (video pixel coordinates) for a job."""
    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    crop = await request.json()
    logger.info("[%s] Received crop region: %s", job_id, crop)

    # Validate required fields
    for key in ("x", "y", "w", "h"):
        if key not in crop:
            raise HTTPException(status_code=400, detail=f"Missing crop field: {key}")

    # Ensure the job directory exists and save crop as JSON
    job_dir = ensure_dir(get_job_dir(job_id))
    crop_path = job_dir / "crop.json"
    with open(crop_path, "w") as f:
        json.dump(crop, f)

    logger.info("[%s] Saved crop.json at %s", job_id, crop_path)
    return {"status": "ok", "crop": crop}


@router.post("/jobs/{job_id}/register-video")
def api_register_video(job_id: str, req: RegisterVideoRequest):
    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    valid, msg = validate_video_path(req.source_video_path)
    if not valid:
        raise HTTPException(status_code=400, detail=msg)

    # Validate video properties
    validation = validate_video(req.source_video_path)
    if not validation["valid"]:
        raise HTTPException(status_code=400, detail=validation.get("error", "Invalid video"))

    update_job(
        job_id,
        source_video_path=req.source_video_path,
        duration_ms=validation.get("duration_ms"),
    )
    return get_job(job_id)


@router.post("/jobs/{job_id}/process")
def api_process_job(job_id: str):
    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    if not job.get("source_video_path"):
        raise HTTPException(status_code=400, detail="No video registered for this job")

    if job.get("status") == "processing":
        raise HTTPException(status_code=409, detail="Job is already processing")

    video_path = job["source_video_path"]
    if video_path != "(demo - synthetic data)":
        valid, msg = validate_video_path(video_path)
        if not valid:
            raise HTTPException(status_code=400, detail=msg)

    update_job_status(job_id, "processing")

    def _run():
        try:
            orch = PipelineOrchestrator()
            orch.run_pipeline(job_id, video_path)
        except Exception as e:
            logger.exception("Pipeline failed for job %s: %s", job_id, e)
            try:
                update_job_status(job_id, "failed")
            except Exception:
                pass
        finally:
            _running_jobs.pop(job_id, None)

    t = threading.Thread(target=_run, daemon=True)
    _running_jobs[job_id] = t
    t.start()

    return {"status": "processing", "job_id": job_id}


@router.post("/jobs/{job_id}/extract-frames")
def api_extract_frames(job_id: str):
    """Extract frames from the video without running the full pipeline.

    Validates, samples, preprocesses, detects relevance + scene changes,
    generates thumbnails, and persists frames to DB. Runs in a background thread.
    """
    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if not job.get("source_video_path"):
        raise HTTPException(status_code=400, detail="No video registered for this job")
    if job.get("status") == "processing":
        raise HTTPException(status_code=409, detail="Job is already processing")

    video_path = job["source_video_path"]
    if video_path != "(demo - synthetic data)":
        valid, msg = validate_video_path(video_path)
        if not valid:
            raise HTTPException(status_code=400, detail=msg)

    update_job_status(job_id, "processing")

    def _run_extract():
        import cv2
        import json as _json
        try:
            logger.info("[%s] Starting frame extraction for %s", job_id, video_path)

            # Load crop region if set
            crop = None
            crop_path = get_job_dir(job_id) / "crop.json"
            if crop_path.is_file():
                with open(crop_path) as cf:
                    crop = _json.load(cf)
                logger.info("[%s] Crop region: x=%d y=%d w=%d h=%d", job_id,
                            crop["x"], crop["y"], crop["w"], crop["h"])

            # 1. Validate
            validation = validate_video(video_path)
            logger.info("[%s] Validation result: valid=%s fps=%s frames=%s",
                        job_id, validation["valid"], validation.get("fps"), validation.get("frame_count"))
            if not validation["valid"]:
                logger.error("[%s] Video validation failed: %s", job_id, validation.get("error"))
                update_job_status(job_id, "failed")
                return

            # 2. Sample frames using OpenCV timestamp-based extraction
            frames_dir = str(ensure_dir(get_job_subdir(job_id, "frames")))
            processed_frames_dir = str(ensure_dir(get_job_subdir(job_id, "processed_frames")))
            from packages.shared.config import FRAME_SAMPLE_INTERVAL_SEC
            sampled = sample_frames(video_path, frames_dir, interval_sec=FRAME_SAMPLE_INTERVAL_SEC)
            logger.info("[%s] Sampled %d frames at %.1fs intervals", job_id, len(sampled), FRAME_SAMPLE_INTERVAL_SEC)
            if not sampled:
                logger.error("[%s] No frames sampled from video", job_id)
                update_job_status(job_id, "failed")
                return

            # 2b. Apply crop region if set — overwrite each frame with cropped version
            if crop:
                cx, cy, cw, ch = crop["x"], crop["y"], crop["w"], crop["h"]
                logger.info("[%s] Cropping %d frames to (%d,%d) %dx%d", job_id, len(sampled), cx, cy, cw, ch)
                for frame in sampled:
                    img = cv2.imread(frame["image_path"])
                    if img is not None:
                        ih, iw = img.shape[:2]
                        # Clamp crop to image bounds
                        x1 = max(0, min(cx, iw))
                        y1 = max(0, min(cy, ih))
                        x2 = max(0, min(cx + cw, iw))
                        y2 = max(0, min(cy + ch, ih))
                        if x2 > x1 and y2 > y1:
                            cropped = img[y1:y2, x1:x2]
                            cv2.imwrite(frame["image_path"], cropped)

            # 3. Preprocess (blur + stability) — read each frame only once
            logger.info("[%s] Preprocessing %d frames...", job_id, len(sampled))
            prev_img = None
            frame_images = {}  # Cache: frame_index -> cv2 image
            for frame in sampled:
                cur_img = cv2.imread(frame["image_path"])
                frame_images[frame["frame_index"]] = cur_img
                if cur_img is not None:
                    frame["blur_score"] = compute_blur_score(cur_img)
                    frame["is_sharp"] = frame["blur_score"] >= 0.15
                    frame["dimensions"] = (cur_img.shape[1], cur_img.shape[0])
                else:
                    frame["blur_score"] = 0.0
                    frame["is_sharp"] = False
                    frame["dimensions"] = (0, 0)

                if prev_img is not None and cur_img is not None:
                    frame["stability_score"] = compute_stability_score(prev_img, cur_img)
                else:
                    frame["stability_score"] = 0.5
                prev_img = cur_img

            # 4. Detect relevance (skip OCR here — full OCR runs later on selected frames)
            logger.info("[%s] Detecting relevance...", job_id)
            for frame in sampled:
                rel = detect_relevance(frame["image_path"], use_ocr=False)
                frame.update(rel)

            # 5. Scene change detection
            logger.info("[%s] Detecting scene changes...", job_id)
            sampled = detect_scene_changes(sampled)

            # 5b. Redundant frame detection — compare scene-change frames
            #     against each other and mark near-duplicates (SSIM > 0.95).
            from packages.core.pipeline.scene_change import _ssim, _gray_resized
            REDUNDANCY_SSIM_THRESHOLD = 0.95
            scene_change_frames = [
                f for f in sampled if f.get("is_scene_change")
            ]
            # Pre-compute grayscale thumbnails using cached images
            _sc_grays = {}
            for f in scene_change_frames:
                img = frame_images.get(f["frame_index"])
                if img is None:
                    img = cv2.imread(f["image_path"])
                if img is not None:
                    _sc_grays[f["frame_index"]] = _gray_resized(img)
            # Compare each scene-change frame to all earlier scene-change frames
            redundant_count = 0
            seen_indices = []
            for f in scene_change_frames:
                fidx = f["frame_index"]
                gray = _sc_grays.get(fidx)
                if gray is None:
                    f["is_redundant"] = False
                    seen_indices.append(fidx)
                    continue
                is_redundant = False
                for prev_idx in seen_indices:
                    prev_gray = _sc_grays.get(prev_idx)
                    if prev_gray is None:
                        continue
                    similarity = _ssim(prev_gray, gray)
                    if similarity > REDUNDANCY_SSIM_THRESHOLD:
                        is_redundant = True
                        break
                f["is_redundant"] = is_redundant
                if is_redundant:
                    redundant_count += 1
                seen_indices.append(fidx)
            # Mark non-scene-change frames as not redundant
            for f in sampled:
                if "is_redundant" not in f:
                    f["is_redundant"] = False
            logger.info("[%s] Redundant frames detected: %d / %d scene-change frames",
                        job_id, redundant_count, len(scene_change_frames))

            # 6. Detect UI elements with bounding boxes + OCR on scene-change frames
            #    Uses batch parallel processing for speed
            logger.info("[%s] Detecting UI elements with bounding boxes...", job_id)
            from packages.core.pipeline.detect_elements import detect_elements_batch
            ocr_frames = [f for f in sampled if f.get("is_scene_change") or f.get("frame_index", 0) == 0]
            if ocr_frames:
                ocr_paths = [f["image_path"] for f in ocr_frames]
                batch_results = detect_elements_batch(ocr_paths, max_workers=4, quick=True)
                for frame, elem_result in zip(ocr_frames, batch_results):
                    frame["element_count"] = elem_result["element_count"]
                    frame["elements"] = elem_result["elements"]
                    frame["annotated_path"] = elem_result.get("annotated_path", "")
                    if elem_result.get("extracted_text"):
                        frame["extracted_text"] = elem_result["extracted_text"]
                    if elem_result.get("ocr_confidence", 0) > 0:
                        frame["ocr_confidence"] = elem_result["ocr_confidence"]

            # 7. Generate thumbnails using cached images (no re-reads)
            logger.info("[%s] Generating thumbnails...", job_id)
            for frame in sampled:
                img = frame_images.get(frame["frame_index"])
                if img is None:
                    img = cv2.imread(frame["image_path"])
                if img is not None:
                    thumb_name = f"frame_thumb_{frame['frame_index']:06d}.jpg"
                    thumb_path = os.path.join(processed_frames_dir, thumb_name)
                    cv2.imwrite(thumb_path, cv2.resize(img, (320, 180)),
                                [cv2.IMWRITE_JPEG_QUALITY, 80])
                    frame["thumbnail_path"] = thumb_path

            # Free cached images to release memory
            frame_images.clear()

            # 8. Clear existing frames for this job (in case of re-extraction)
            existing_frames = get_frames_for_job(job_id)
            if existing_frames:
                from packages.shared.database import _get_conn
                conn = _get_conn()
                conn.execute("DELETE FROM frames WHERE job_id = ?", (job_id,))
                conn.commit()
                logger.info("[%s] Cleared %d old frames", job_id, len(existing_frames))

            # 9. Persist frames to DB
            logger.info("[%s] Persisting %d frames to DB...", job_id, len(sampled))
            for frame in sampled:
                create_frame(
                    frame_id=generate_id(),
                    job_id=job_id,
                    frame_index=frame["frame_index"],
                    timestamp_ms=frame["timestamp_ms"],
                    source_image_path=frame["image_path"],
                    blur_score=frame.get("blur_score", 0),
                    stability_score=frame.get("stability_score", 0),
                    relevance_score=frame.get("relevance_score", 0),
                    matched_keywords=frame.get("matched_keywords", []),
                    extracted_text=frame.get("extracted_text", ""),
                    ocr_confidence=frame.get("ocr_confidence", 0),
                    candidate_score=0.0 if frame.get("is_redundant") else frame.get("visual_importance", frame.get("candidate_score", 0)),
                )

            update_job_status(job_id, "completed", frame_count=len(sampled))

            # Auto-generate a better title if current title is generic
            current_job = get_job(job_id)
            current_title = current_job.get("title", "") if current_job else ""
            if current_title.startswith("Recording "):
                auto_title = _generate_auto_title(job_id)
                if auto_title:
                    update_job(job_id, title=auto_title)
                    logger.info("[%s] Auto-titled job: %s", job_id, auto_title)

            logger.info("[%s] Frame extraction DONE: %d frames", job_id, len(sampled))

        except Exception as e:
            import traceback
            err_msg = traceback.format_exc()
            logger.exception("[%s] Frame extraction FAILED: %s", job_id, e)
            # Write error to file for debugging
            err_path = os.path.join(str(get_job_subdir(job_id, ".")), "extraction_error.txt")
            try:
                with open(err_path, "w") as ef:
                    ef.write(err_msg)
            except Exception:
                pass
            try:
                update_job_status(job_id, "failed")
            except Exception:
                pass
        finally:
            _running_jobs.pop(job_id, None)

    t = threading.Thread(target=_run_extract, daemon=True)
    _running_jobs[job_id] = t
    t.start()

    return {"status": "processing", "job_id": job_id}


@router.get("/jobs/{job_id}/status")
def api_job_status(job_id: str):
    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return {
        "job_id": job_id,
        "status": job["status"],
        "frame_count": job.get("frame_count"),
        "screenshot_count": job.get("screenshot_count"),
        "is_running": job_id in _running_jobs,
    }


# --- Frame endpoints ---

@router.get("/jobs/{job_id}/frames")
def api_get_frames(
    job_id: str,
    min_relevance: float = Query(0.0, ge=0.0, le=1.0),
    limit: int = Query(500, ge=1, le=2000),
):
    _validate_job_id(job_id)
    return get_frames_for_job(job_id, min_relevance=min_relevance, limit=limit)


@router.get("/jobs/{job_id}/video-info")
def api_video_info(job_id: str):
    """Return video duration, fps, and dimensions using OpenCV."""
    import cv2

    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if not job.get("source_video_path"):
        raise HTTPException(status_code=400, detail="No video registered for this job")

    video_path = job["source_video_path"]
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        raise HTTPException(status_code=500, detail="Failed to open video file")

    try:
        fps = cap.get(cv2.CAP_PROP_FPS)
        frame_count = cap.get(cv2.CAP_PROP_FRAME_COUNT)
        width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
        height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))

        # WebM files report bogus fps=1000 — use seek-to-end fallback
        if fps >= 500:
            cap.set(cv2.CAP_PROP_POS_AVI_RATIO, 1)
            duration_ms = int(cap.get(cv2.CAP_PROP_POS_MSEC))
            if duration_ms <= 0:
                # Read through to find last timestamp
                cap.set(cv2.CAP_PROP_POS_MSEC, 0)
                last_ms = 0
                while True:
                    ret = cap.grab()
                    if not ret:
                        break
                    last_ms = cap.get(cv2.CAP_PROP_POS_MSEC)
                duration_ms = int(last_ms)
        else:
            duration_ms = int((frame_count / fps) * 1000) if fps > 0 else 0
    finally:
        cap.release()

    return {
        "duration_ms": duration_ms,
        "fps": round(fps, 2),
        "width": width,
        "height": height,
    }


@router.post("/jobs/{job_id}/extract-frame-at")
async def api_extract_frame_at(job_id: str, request: Request):
    """Extract a single frame from the video at a given timestamp (ms)."""
    import cv2

    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    if not job.get("source_video_path"):
        raise HTTPException(status_code=400, detail="No video registered for this job")

    data = await request.json()
    timestamp_ms = data.get("timestamp_ms")
    if timestamp_ms is None:
        raise HTTPException(status_code=400, detail="Missing required field: timestamp_ms")
    timestamp_ms = int(timestamp_ms)

    video_path = job["source_video_path"]
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        raise HTTPException(status_code=500, detail="Failed to open video file")

    try:
        cap.set(cv2.CAP_PROP_POS_MSEC, timestamp_ms)
        ret, frame_img = cap.read()
    finally:
        cap.release()

    if not ret or frame_img is None:
        raise HTTPException(status_code=400, detail=f"Could not read frame at {timestamp_ms}ms")

    # Apply crop if crop.json exists
    crop_path = get_job_dir(job_id) / "crop.json"
    if crop_path.is_file():
        with open(crop_path) as cf:
            crop = json.load(cf)
        cx, cy, cw, ch = crop["x"], crop["y"], crop["w"], crop["h"]
        ih, iw = frame_img.shape[:2]
        x1 = max(0, min(cx, iw))
        y1 = max(0, min(cy, ih))
        x2 = max(0, min(cx + cw, iw))
        y2 = max(0, min(cy + ch, ih))
        if x2 > x1 and y2 > y1:
            frame_img = frame_img[y1:y2, x1:x2]

    # Save frame
    frames_dir = str(ensure_dir(get_job_subdir(job_id, "frames")))
    frame_filename = f"manual_frame_{timestamp_ms}ms.png"
    frame_path = os.path.join(frames_dir, frame_filename)
    cv2.imwrite(frame_path, frame_img)

    # Generate thumbnail
    processed_frames_dir = str(ensure_dir(get_job_subdir(job_id, "processed_frames")))
    thumb_filename = f"frame_thumb_manual_{timestamp_ms}ms.jpg"
    thumb_path = os.path.join(processed_frames_dir, thumb_filename)
    cv2.imwrite(thumb_path, cv2.resize(frame_img, (320, 180)),
                [cv2.IMWRITE_JPEG_QUALITY, 80])

    # Determine frame_index: use 100000 + count of existing manual frames
    existing_frames = get_frames_for_job(job_id)
    manual_count = sum(1 for f in existing_frames if f.get("frame_index", 0) >= 100000)
    frame_index = 100000 + manual_count

    # Create frame record in DB
    frame_id = generate_id()
    create_frame(
        frame_id=frame_id,
        job_id=job_id,
        frame_index=frame_index,
        timestamp_ms=timestamp_ms,
        source_image_path=frame_path,
        blur_score=0.0,
        stability_score=0.0,
        relevance_score=0.0,
        matched_keywords=[],
        extracted_text="",
        ocr_confidence=0.0,
        candidate_score=0.0,
    )

    return get_frame(frame_id)


# --- Screenshot endpoints ---

@router.get("/jobs/{job_id}/screenshots")
def api_get_screenshots(
    job_id: str,
    section_type: str | None = Query(None),
    accepted_only: bool = Query(False),
):
    _validate_job_id(job_id)
    return get_screenshots_for_job(job_id, section_type=section_type, accepted_only=accepted_only)


@router.patch("/screenshots/{screenshot_id}")
def api_update_screenshot(screenshot_id: str, req: UpdateScreenshotRequest):
    existing = get_screenshot(screenshot_id)
    if not existing:
        raise HTTPException(status_code=404, detail="Screenshot not found")

    updates = {}
    if req.accepted is not None:
        updates["accepted"] = req.accepted
    if req.notes is not None:
        updates["notes"] = req.notes
    if req.section_type is not None:
        updates["section_type"] = req.section_type.value
    if req.order_index is not None:
        updates["order_index"] = req.order_index

    if updates:
        update_screenshot(screenshot_id, **updates)

    return get_screenshot(screenshot_id)


# --- Promote frame to screenshot (manual selection) ---

@router.post("/frames/{frame_id}/promote")
def api_promote_frame(frame_id: str):
    """Promote a raw frame to a screenshot (manual selection)."""
    frame = get_frame(frame_id)
    if not frame:
        raise HTTPException(status_code=404, detail="Frame not found")

    job_id = frame["job_id"]
    _validate_job_id(job_id)

    # Check if already promoted
    existing = get_screenshots_for_job(job_id)
    for ss in existing:
        if ss.get("source_frame_id") == f"frame_{frame['frame_index']:06d}":
            return ss  # Already a screenshot

    # Copy frame to screenshots dir
    screenshots_dir = ensure_dir(get_job_subdir(job_id, "screenshots"))
    thumbnails_dir = ensure_dir(get_job_subdir(job_id, "thumbnails"))

    idx = len(existing)
    dst = str(screenshots_dir / f"screenshot_{idx:03d}.png")
    thumb = str(thumbnails_dir / f"thumb_{idx:03d}.png")

    import cv2
    img = cv2.imread(frame["source_image_path"])
    if img is not None:
        cv2.imwrite(dst, img)
        cv2.imwrite(thumb, cv2.resize(img, (320, 180)))

    screenshot_id = generate_id()
    create_screenshot(
        screenshot_id=screenshot_id,
        job_id=job_id,
        source_frame_id=f"frame_{frame['frame_index']:06d}",
        image_path=dst,
        thumbnail_path=thumb,
        captured_at_ms=frame.get("timestamp_ms", 0),
        section_type="unknown",
        confidence=frame.get("candidate_score", 0),
        rationale="Manually selected",
        matched_keywords=json.dumps(frame.get("matched_keywords", [])),
        extracted_text=frame.get("extracted_text", ""),
    )

    return get_screenshot(screenshot_id)


@router.delete("/screenshots/{screenshot_id}")
def api_delete_screenshot(screenshot_id: str):
    """Remove a screenshot (manual deselection)."""
    existing = get_screenshot(screenshot_id)
    if not existing:
        raise HTTPException(status_code=404, detail="Screenshot not found")

    from packages.shared.database import _get_conn
    conn = _get_conn()
    conn.execute("DELETE FROM sections WHERE screenshot_id = ?", (screenshot_id,))
    conn.execute("DELETE FROM screenshots WHERE id = ?", (screenshot_id,))
    conn.commit()
    return {"deleted": screenshot_id}


@router.post("/jobs/{job_id}/select-all")
def api_select_all(job_id: str):
    """Promote all non-selected frames to screenshots (select all)."""
    import cv2

    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    frames = get_frames_for_job(job_id, limit=2000)
    existing = get_screenshots_for_job(job_id)
    already_selected = {ss.get("source_frame_id") for ss in existing}

    screenshots_dir = ensure_dir(get_job_subdir(job_id, "screenshots"))
    thumbnails_dir = ensure_dir(get_job_subdir(job_id, "thumbnails"))

    added = 0
    idx = len(existing)
    for frame in frames:
        ref = f"frame_{frame['frame_index']:06d}"
        if ref in already_selected:
            continue
        dst = str(screenshots_dir / f"screenshot_{idx:03d}.png")
        thumb = str(thumbnails_dir / f"thumb_{idx:03d}.png")
        img = cv2.imread(frame["source_image_path"])
        if img is not None:
            cv2.imwrite(dst, img)
            cv2.imwrite(thumb, cv2.resize(img, (320, 180)))
        screenshot_id = generate_id()
        create_screenshot(
            screenshot_id=screenshot_id,
            job_id=job_id,
            source_frame_id=ref,
            image_path=dst,
            thumbnail_path=thumb,
            captured_at_ms=frame.get("timestamp_ms", 0),
            section_type="unknown",
            confidence=frame.get("candidate_score", 0),
            rationale="Manually selected",
            matched_keywords=json.dumps(frame.get("matched_keywords", [])),
            extracted_text=frame.get("extracted_text", ""),
        )
        added += 1
        idx += 1

    return {"added": added, "total_selected": len(existing) + added}


@router.post("/jobs/{job_id}/unselect-all")
def api_unselect_all(job_id: str):
    """Remove all screenshots for this job (unselect all)."""
    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    from packages.shared.database import _get_conn
    conn = _get_conn()
    conn.execute("DELETE FROM sections WHERE job_id = ?", (job_id,))
    conn.execute("DELETE FROM screenshots WHERE job_id = ?", (job_id,))
    conn.commit()
    return {"status": "cleared"}


# --- Section endpoints ---

@router.get("/jobs/{job_id}/sections")
def api_get_sections(job_id: str):
    _validate_job_id(job_id)
    return get_sections_for_job(job_id)


# --- Report endpoints ---

@router.post("/jobs/{job_id}/report")
def api_generate_report(job_id: str):
    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    screenshots = get_screenshots_for_job(job_id, accepted_only=True)
    sections = get_sections_for_job(job_id)

    if not screenshots and not sections:
        raise HTTPException(status_code=400, detail="No screenshots or sections to include in report")

    reports_dir = ensure_dir(get_job_subdir(job_id, "reports"))
    html_path = str(reports_dir / f"report_{job_id}.html")
    pdf_path = str(reports_dir / f"report_{job_id}.pdf")

    job_meta = {
        "job_id": job_id,
        "video_path": job.get("source_video_path", "unknown"),
        "status": job.get("status", "unknown"),
    }

    # Build section/screenshot pairs
    section_dicts = []
    screenshot_dicts = []
    for i, ss in enumerate(screenshots):
        # Find matching section
        matching_section = next(
            (s for s in sections if s.get("screenshot_id") == ss["id"]), None
        )
        if matching_section:
            section_dicts.append(matching_section)
        else:
            section_dicts.append({
                "heading": ss.get("section_type", "Section").replace("_", " ").title(),
                "section_type": ss.get("section_type", "unknown"),
                "summary": ss.get("extracted_text", "")[:200],
                "key_points": ss.get("matched_keywords", []),
                "order_suggestion": i,
            })
        screenshot_dicts.append(ss)

    generate_html_report(job_meta, section_dicts, screenshot_dicts, html_path)
    generate_pdf_report(html_path, pdf_path)

    # Check if report already exists
    existing = get_report_for_job(job_id)
    if not existing:
        create_report(
            report_id=generate_id(),
            job_id=job_id,
            html_path=html_path,
            pdf_path=pdf_path,
        )

    return get_report_for_job(job_id)


@router.get("/jobs/{job_id}/report")
def api_get_report(job_id: str):
    _validate_job_id(job_id)
    report = get_report_for_job(job_id)
    if not report:
        raise HTTPException(status_code=404, detail="Report not found. Generate one first.")
    return report


@router.get("/jobs/{job_id}/report/html")
def api_get_report_html(job_id: str):
    _validate_job_id(job_id)
    report = get_report_for_job(job_id)
    if not report or not report.get("html_path"):
        raise HTTPException(status_code=404, detail="HTML report not found")
    html_path = report["html_path"]
    if not os.path.isfile(html_path):
        raise HTTPException(status_code=404, detail="Report file missing from disk")
    with open(html_path, "r", encoding="utf-8") as f:
        content = f.read()
    return HTMLResponse(content=content)


# --- Artifact serving ---

@router.get("/artifacts/{job_id}/{artifact_type}/{filename}")
def api_serve_artifact(job_id: str, artifact_type: str, filename: str):
    _validate_job_id(job_id)

    if artifact_type not in _VALID_ARTIFACT_TYPES:
        raise HTTPException(status_code=400, detail=f"Invalid artifact type: {artifact_type}")

    # Sanitize filename to prevent traversal
    if ".." in filename or "/" in filename or "\\" in filename:
        raise HTTPException(status_code=400, detail="Invalid filename")

    artifact_dir = get_job_subdir(job_id, artifact_type)
    file_path = (artifact_dir / filename).resolve()

    # Ensure path is within the job directory
    if not str(file_path).startswith(str(JOBS_DIR.resolve())):
        raise HTTPException(status_code=403, detail="Access denied")

    if not file_path.is_file():
        raise HTTPException(status_code=404, detail="Artifact not found")

    return FileResponse(str(file_path))


# --- Demo ---

@router.post("/demo/seed")
def api_seed_demo():
    """Seed demo data for testing the UI without a real video."""
    import numpy as np
    try:
        import cv2
    except ImportError:
        raise HTTPException(status_code=500, detail="OpenCV not installed")

    DEMO_SECTIONS = [
        {"section_type": "demographics", "heading": "Applicant Demographics",
         "summary": "Personal information section showing applicant details.",
         "key_points": ["Name: Jane Doe", "DOB: 1985-03-15", "Address: 123 Main St"],
         "keywords": ["name", "date of birth", "address"],
         "color": (200, 220, 240),
         "text_lines": ["APPLICANT DEMOGRAPHICS", "", "Name: Jane Doe", "Date of Birth: 03/15/1985",
                         "Address: 123 Main St", "City: Springfield, IL 62701"]},
        {"section_type": "income", "heading": "Income Verification",
         "summary": "Income details showing monthly earnings and employment status.",
         "key_points": ["Monthly Income: $2,450", "Employer: ABC Corp", "Full-time"],
         "keywords": ["income", "wages", "salary", "employment"],
         "color": (220, 240, 200),
         "text_lines": ["INCOME VERIFICATION", "", "Employment: Full-time", "Employer: ABC Corp",
                         "Monthly Gross: $2,450.00", "Annual: $29,400.00"]},
        {"section_type": "household", "heading": "Household Composition",
         "summary": "Household members listing showing family size.",
         "key_points": ["Household Size: 3", "Dependents: 1 child", "Spouse: John Doe"],
         "keywords": ["household", "members", "dependents", "family"],
         "color": (240, 220, 200),
         "text_lines": ["HOUSEHOLD COMPOSITION", "", "Total Members: 3", "",
                         "1. Jane Doe (Applicant) - Age 40", "2. John Doe (Spouse) - Age 42",
                         "3. Emily Doe (Child) - Age 8"]},
        {"section_type": "eligibility", "heading": "Eligibility Determination",
         "summary": "Eligibility screening results for benefits programs.",
         "key_points": ["Status: Potentially Eligible", "Program: Medicaid", "FPL: 138%"],
         "keywords": ["eligible", "eligibility", "qualify"],
         "color": (200, 240, 220),
         "text_lines": ["ELIGIBILITY DETERMINATION", "", "Program: Medicaid",
                         "Federal Poverty Level: 138%", "Status: POTENTIALLY ELIGIBLE"]},
        {"section_type": "table", "heading": "Benefits Summary Table",
         "summary": "Summary table showing available benefits programs.",
         "key_points": ["Medicaid: Eligible", "SNAP: Review Required", "CHIP: Eligible"],
         "keywords": ["table", "amount", "benefits"],
         "color": (230, 230, 240),
         "text_lines": ["BENEFITS SUMMARY", "", "Program     | Status    | Amount",
                         "------------|-----------|------", "Medicaid    | Eligible  | Full",
                         "SNAP        | Review    | $234/mo", "CHIP        | Eligible  | Full"]},
        {"section_type": "policy_guidance", "heading": "Policy Guidance Notes",
         "summary": "Relevant policy guidance for eligibility determination.",
         "key_points": ["42 CFR 435.603", "5% FPL income disregard", "Retroactive: 3 months"],
         "keywords": ["policy", "regulation", "guidance"],
         "color": (240, 230, 220),
         "text_lines": ["POLICY GUIDANCE", "", "Reference: 42 CFR 435.603",
                         "MAGI-Based Income Methodology", "",
                         "- 5% FPL income disregard applies", "- Retroactive coverage up to 3 months"]},
    ]

    job_id = generate_id()
    screenshots_dir = ensure_dir(get_job_subdir(job_id, "screenshots"))
    thumbnails_dir = ensure_dir(get_job_subdir(job_id, "thumbnails"))
    ensure_dir(get_job_subdir(job_id, "input"))

    create_job(job_id=job_id, title="Demo: Benefits Eligibility Review",
               source_video_path="(demo - synthetic data)")

    for i, sec in enumerate(DEMO_SECTIONS):
        screenshot_id = generate_id()
        ts_ms = (i + 1) * 5000

        # Create synthetic image
        img = np.full((720, 1280, 3), sec["color"], dtype=np.uint8)
        cv2.rectangle(img, (0, 0), (1280, 60), (50, 70, 90), -1)
        cv2.putText(img, "PolicyCapture Local - Demo", (20, 40),
                    cv2.FONT_HERSHEY_SIMPLEX, 0.8, (255, 255, 255), 2)
        cv2.rectangle(img, (40, 80), (1240, 680), (180, 180, 180), 2)
        y = 130
        for line in sec["text_lines"]:
            if not line:
                y += 20
                continue
            scale = 0.85 if line.isupper() else 0.7
            thick = 2 if line.isupper() else 1
            color = (20, 50, 100) if line.isupper() else (30, 30, 30)
            cv2.putText(img, line, (70, y), cv2.FONT_HERSHEY_SIMPLEX, scale, color, thick)
            y += 35

        img_path = str(screenshots_dir / f"screenshot_{i:03d}.png")
        thumb_path = str(thumbnails_dir / f"thumb_{i:03d}.png")
        cv2.imwrite(img_path, img)
        cv2.imwrite(thumb_path, cv2.resize(img, (320, 180)))

        create_screenshot(
            screenshot_id=screenshot_id, job_id=job_id,
            source_frame_id=f"frame_{i:06d}",
            image_path=img_path, thumbnail_path=thumb_path,
            captured_at_ms=ts_ms, section_type=sec["section_type"],
            confidence=0.85 + (i * 0.02),
            rationale=f"Matched: {', '.join(sec['keywords'][:3])}",
            matched_keywords=json.dumps(sec["keywords"]),
            extracted_text=" ".join(sec["text_lines"]),
        )

        create_section(
            section_id=generate_id(), job_id=job_id, screenshot_id=screenshot_id,
            heading=sec["heading"], section_type=sec["section_type"],
            summary=sec["summary"],
            key_points=json.dumps(sec["key_points"]),
            confidence=0.85 + (i * 0.02), final_order=i,
        )

    update_job_status(job_id, "completed", screenshot_count=len(DEMO_SECTIONS))

    return {"job_id": job_id, "message": "Demo data seeded", "screenshot_count": len(DEMO_SECTIONS)}


# ==========================================================================
#  OCR, Entity Extraction & Search
# ==========================================================================

@router.post("/jobs/{job_id}/run-ocr")
def api_run_ocr(job_id: str, force: bool = Query(False)):
    """Run OCR + entity extraction on all selected screenshots for a job.

    Uses parallel batch processing for speed — processes multiple screenshots
    concurrently with multi-strategy adaptive OCR.

    Set force=true to re-run OCR on all frames, even those already processed.
    """
    _validate_job_id(job_id)
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    screenshots_list = get_screenshots_for_job(job_id)
    if not screenshots_list:
        raise HTTPException(status_code=400, detail="No screenshots to process")

    from packages.core.pipeline.detect_elements import detect_elements_batch, HAS_TESSERACT
    try:
        from packages.core.pipeline.extract_entities import extract_entities
    except ImportError:
        extract_entities = None

    # Build list of screenshots that need processing and resolve their image paths
    frames = get_frames_for_job(job_id)
    frame_map = {}
    for f in frames:
        frame_map[f["frame_index"]] = f

    to_process = []  # (screenshot, image_path, frame_index)
    for ss in screenshots_list:
        if not force and ss.get("extracted_text", ""):
            continue

        frame_ref = ss.get("source_frame_id", "")
        frame_index = -1
        if frame_ref.startswith("frame_"):
            try:
                frame_index = int(frame_ref.replace("frame_", ""))
            except ValueError:
                pass

        image_path = ss.get("image_path", "")
        if not image_path or not os.path.isfile(image_path):
            f = frame_map.get(frame_index)
            if f:
                image_path = f["source_image_path"]

        if image_path and os.path.isfile(image_path):
            to_process.append((ss, image_path, frame_index))

    if not to_process:
        return {
            "processed": 0,
            "total": len(screenshots_list),
            "ocr_available": HAS_TESSERACT,
            "results": [],
        }

    # Run batch OCR in parallel (full quality, not quick mode)
    image_paths = [item[1] for item in to_process]
    batch_results = detect_elements_batch(image_paths, max_workers=4, quick=False)

    processed = 0
    results = []

    for (ss, image_path, frame_index), elem_result in zip(to_process, batch_results):
        text = elem_result.get("extracted_text", "")
        elements = elem_result.get("elements", [])
        tables = elem_result.get("tables", [])
        checkboxes = elem_result.get("checkboxes", [])
        form_fields = elem_result.get("form_fields", [])
        structured_data = elem_result.get("structured_data", {})
        confidence = elem_result.get("ocr_confidence", 0.0)
        if not confidence and elements:
            confs = [e["confidence"] for e in elements if e["confidence"] > 0]
            confidence = sum(confs) / len(confs) if confs else 0.0

        # Run NER entity extraction (now includes form_data, lists, section_headers)
        entities = {}
        if extract_entities and text:
            entities = extract_entities(text)

        # Include visual detections in the entity data payload
        if tables:
            entities["tables"] = tables
        if checkboxes:
            entities["checkboxes"] = checkboxes
        if form_fields:
            entities["form_fields"] = form_fields
        if structured_data:
            # Merge text-extracted structured data with NER form_data
            if "form_data" in entities and structured_data.get("key_value_pairs"):
                # Visual structured data supplements NER-extracted form_data
                entities["structured_data_visual"] = structured_data
            elif structured_data.get("key_value_pairs") or structured_data.get("lists"):
                entities["structured_data_visual"] = structured_data

        # Store entities+tables+forms as JSON in matched_keywords field
        entity_data = json.dumps(entities) if entities else "[]"

        # Update screenshot with OCR text and entities
        update_screenshot(ss["id"], extracted_text=text, matched_keywords=entity_data, notes=ss.get("notes", ""))

        # Also update the source frame if it exists
        if frame_index >= 0:
            f = frame_map.get(frame_index)
            if f:
                from packages.shared.database import _get_conn
                conn = _get_conn()
                conn.execute(
                    "UPDATE frames SET extracted_text = ?, ocr_confidence = ? WHERE id = ?",
                    (text, confidence, f["id"]),
                )
                conn.commit()

        processed += 1
        results.append({
            "screenshot_id": ss["id"],
            "text_length": len(text),
            "entity_count": entities.get("summary", {}).get("total_entities", 0) if entities else 0,
            "table_count": len(tables),
            "ocr_confidence": round(confidence, 3),
            "ocr_strategy": elem_result.get("ocr_strategy", "unknown"),
        })

    return {
        "processed": processed,
        "total": len(screenshots_list),
        "ocr_available": HAS_TESSERACT,
        "results": results,
    }


@router.get("/jobs/{job_id}/ocr-data")
def api_get_ocr_data(job_id: str):
    """Get all OCR text and entities for a job's screenshots."""
    _validate_job_id(job_id)

    screenshots_list = get_screenshots_for_job(job_id)
    frames = get_frames_for_job(job_id)
    frame_map = {f"frame_{f['frame_index']:06d}": f for f in frames}

    items = []
    for ss in screenshots_list:
        ref = ss.get("source_frame_id", "")
        frame = frame_map.get(ref)

        text = ss.get("extracted_text", "") or (frame.get("extracted_text", "") if frame else "")

        # Parse entities from matched_keywords
        entities = {}
        kw = ss.get("matched_keywords", "[]")
        if kw and kw != "[]":
            try:
                entities = json.loads(kw) if isinstance(kw, str) else kw
            except (json.JSONDecodeError, TypeError):
                entities = {}

        # Get image URLs
        image_path = ss.get("image_path", "")
        thumb_path = ss.get("thumbnail_path", "")

        items.append({
            "screenshot_id": ss["id"],
            "frame_ref": ref,
            "frame_index": frame["frame_index"] if frame else -1,
            "timestamp_ms": ss.get("captured_at_ms", frame.get("timestamp_ms", 0) if frame else 0),
            "image_path": image_path,
            "thumbnail_path": thumb_path,
            "extracted_text": text,
            "entities": entities,
            "ocr_confidence": frame.get("ocr_confidence", 0) if frame else 0,
            "notes": ss.get("notes", ""),
            "rationale": ss.get("rationale", ""),
            "section_type": ss.get("section_type", "unknown"),
        })

    items.sort(key=lambda x: x["timestamp_ms"])
    return items


@router.get("/jobs/{job_id}/search")
def api_search_text(job_id: str, q: str = Query("", min_length=1)):
    """Search across all extracted text for a job."""
    _validate_job_id(job_id)

    query = q.lower().strip()
    if not query:
        return []

    frames = get_frames_for_job(job_id)
    screenshots_list = get_screenshots_for_job(job_id)

    # Build text index from both frames and screenshots
    results = []
    seen_refs = set()

    for ss in screenshots_list:
        ref = ss.get("source_frame_id", "")
        text = ss.get("extracted_text", "")
        if not text:
            continue
        if query in text.lower():
            # Find matching snippets
            snippets = _find_snippets(text, query)
            results.append({
                "type": "screenshot",
                "id": ss["id"],
                "frame_ref": ref,
                "timestamp_ms": ss.get("captured_at_ms", 0),
                "text": text,
                "snippets": snippets,
                "match_count": text.lower().count(query),
                "notes": ss.get("notes", ""),
            })
            seen_refs.add(ref)

    for f in frames:
        ref = f"frame_{f['frame_index']:06d}"
        if ref in seen_refs:
            continue
        text = f.get("extracted_text", "")
        if not text:
            continue
        if query in text.lower():
            snippets = _find_snippets(text, query)
            results.append({
                "type": "frame",
                "id": f["id"],
                "frame_ref": ref,
                "timestamp_ms": f.get("timestamp_ms", 0),
                "text": text,
                "snippets": snippets,
                "match_count": text.lower().count(query),
            })

    results.sort(key=lambda x: x["match_count"], reverse=True)
    return results


@router.post("/jobs/{job_id}/export-docx")
def api_export_docx(job_id: str, indices: str = Query("")):
    """Export OCR data as a .docx Word document."""
    _validate_job_id(job_id)
    selected_indices: set[int] | None = None
    if indices.strip():
        try:
            selected_indices = {int(i) for i in indices.split(",") if i.strip().isdigit()}
        except ValueError:
            pass
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    frames = get_frames_for_job(job_id)
    screenshots_list = get_screenshots_for_job(job_id)

    # Build frame data (same logic as ocr-data endpoint)
    ss_by_frame: dict[str, dict] = {}
    for ss in screenshots_list:
        ref = ss.get("source_frame_id", "")
        if ref:
            ss_by_frame[ref] = ss

    items = []
    for f in frames:
        ref = f"frame_{f['frame_index']:06d}"
        ss = ss_by_frame.get(ref, {})
        text = ss.get("extracted_text") or f.get("extracted_text", "")
        conf = f.get("ocr_confidence", 0)
        ts = f.get("timestamp_ms", 0)
        entities_raw = ss.get("entities") or f.get("entities")
        entities = json.loads(entities_raw) if isinstance(entities_raw, str) else (entities_raw or {})
        notes = ss.get("notes", "")
        items.append({
            "timestamp_ms": ts,
            "text": text or "",
            "confidence": conf,
            "entities": entities,
            "notes": notes,
        })
    items.sort(key=lambda x: x["timestamp_ms"])

    # Filter to selected indices if provided
    if selected_indices is not None:
        items = [item for i, item in enumerate(items) if i in selected_indices]

    # Build DOCX
    doc = Document()
    title = doc.add_heading(job.get("title", "OCR Export"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for item in items:
        ts_s = item["timestamp_ms"] // 1000
        ts_m = ts_s // 60
        ts_sec = ts_s % 60
        time_str = f"{ts_m:02d}:{ts_sec:02d}"
        conf = item["confidence"]

        heading = f"Frame at {time_str}"
        if conf > 0:
            heading += f" (confidence: {round(conf)}%)"
        doc.add_heading(heading, level=2)

        # Extracted text
        if item["text"]:
            p = doc.add_paragraph(item["text"])
            p.style.font.size = Pt(10)
        else:
            p = doc.add_paragraph("[No text extracted]")
            p.runs[0].italic = True
            p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Entities
        ent = item["entities"]
        cats = ent.get("categories", {})
        has_entities = any(v for v in cats.values() if v)
        if has_entities:
            doc.add_heading("Entities", level=3)
            for etype, vals in cats.items():
                if vals:
                    label = etype.replace("_", " ").title()
                    p = doc.add_paragraph()
                    run = p.add_run(f"{label}: ")
                    run.bold = True
                    p.add_run(", ".join(vals))

        # Form data
        form_data = ent.get("form_data", {})
        kv_pairs = form_data.get("key_value_pairs", [])
        if kv_pairs:
            doc.add_heading("Form Fields", level=3)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Field"
            hdr[1].text = "Value"
            for kv in kv_pairs:
                row = table.add_row().cells
                row[0].text = kv.get("key", "")
                row[1].text = kv.get("value", "")

        # Tables
        tables = ent.get("tables", [])
        for ti, tbl in enumerate(tables):
            doc.add_heading(f"Table {ti + 1}", level=3)
            rows = tbl.get("rows", [])
            if rows:
                t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                t.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        t.rows[ri].cells[ci].text = cell or ""

        # Notes
        if item["notes"]:
            doc.add_heading("Notes", level=3)
            doc.add_paragraph(item["notes"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"ocr_export_{job_id[:8]}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _find_snippets(text, query, context=60):
    """Extract text snippets around query matches."""
    lower = text.lower()
    snippets = []
    start = 0
    while True:
        idx = lower.find(query, start)
        if idx == -1:
            break
        s = max(0, idx - context)
        e = min(len(text), idx + len(query) + context)
        prefix = "..." if s > 0 else ""
        suffix = "..." if e < len(text) else ""
        snippet = prefix + text[s:e] + suffix
        snippets.append({
            "text": snippet,
            "match_start": idx - s + len(prefix),
            "match_end": idx - s + len(prefix) + len(query),
        })
        start = idx + len(query)
        if len(snippets) >= 5:
            break
    return snippets


@router.put("/screenshots/{screenshot_id}/notes")
async def api_update_notes(screenshot_id: str, request: Request):
    """Update notes/annotations for a screenshot."""
    existing = get_screenshot(screenshot_id)
    if not existing:
        raise HTTPException(status_code=404, detail="Screenshot not found")

    data = await request.json()
    notes = data.get("notes", "")
    update_screenshot(screenshot_id, notes=notes)
    return get_screenshot(screenshot_id)
